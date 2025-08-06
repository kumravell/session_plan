from flask import Flask, request, render_template, redirect, url_for,abort, send_from_directory
from openpyxl import load_workbook, Workbook
import os
from io import BytesIO
from docx import Document
from flask import jsonify
import html
from bs4 import BeautifulSoup
import json
import docx
import flask
import werkzeug
from flask import jsonify
import mysql.connector
import psycopg2
from flask import Flask, session, request, redirect, url_for, render_template


app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "fallback-secret")

# app.secret_key = os.urandom(24)
FILES_DIR = 'files'

DB_CONFIG = {
    'dbname': os.environ.get('DB_NAME'),
    'user': os.environ.get('DB_USER'),
    'password': os.environ.get('DB_PASSWORD'),
    'host': os.environ.get('DB_HOST'),
    'port': os.environ.get('DB_PORT')
}

# DB_CONFIG = {
#     'dbname': 'siesadmin',
#     'user': 'postgres',
#     'password': 'admin',
#     'host': 'localhost',
#     'port': '5432'
# }

user_credentials = {
    'username': 'siesadmin',
    'password': 'sies400706'
}



@app.route('/login', methods=['POST'])
def login():
    username = request.form['username']
    password = request.form['password']
    if username == user_credentials['username'] and password == user_credentials['password']:
        session['logged_in'] = True  # Set session to True when logged in
        return redirect(url_for('dates', table='general_holidays'))  # Redirect to the dates page
    else:
        return "Invalid username or password"

@app.route('/logout')
def logout():
    session.clear()  # Clear all session data
    return redirect(url_for('index'))  # Redirect to the index or login page


@app.route('/dates')
def dates():
    if 'logged_in' not in session:  # Check if the user is logged in
        return redirect(url_for('index'))  # Redirect to login if not logged in

    table = request.args.get('table', 'general_holidays')  # Get the table name from the URL
    try:
        connection = psycopg2.connect(**DB_CONFIG)
        cursor = connection.cursor()
        cursor.execute(f"SELECT id, dates FROM {table}")
        rows = cursor.fetchall()
        cursor.close()
        connection.close()

        holidays = [{"id": row[0], "date": row[1].strftime('%Y-%m-%d')} for row in rows]
        return render_template('dates.html', holidays=holidays, table=table)
    
    except psycopg2.Error as err:
        return f"Error fetching data: {err}"







# @app.route('/dates')
# def dates():
#     holidays = Holiday.query.all() 
#     return render_template('dates.html', holidays=holidays)



# @app.route('/delete_holiday', methods=['POST'])
# def delete_holiday():
#     data = request.get_json()
#     holiday_id = data['holiday_id']
#     holiday = Holiday.query.get(holiday_id)
#     if holiday:
#         db.session.delete(holiday)
#         db.session.commit()
#         return jsonify({'success': True})
#     return jsonify({'success': False, 'message': 'Holiday not found'}), 404




@app.route('/fetch_holidays', methods=['GET'])
def fetch_holidays():
    table = request.args.get('table', 'general_holidays')  # Default to 'general_holidays' if not specified
    try:
        connection = psycopg2.connect(**DB_CONFIG)
        cursor = connection.cursor()
        cursor.execute(f"SELECT id, dates FROM {table}")  # Ensure you are fetching the correct column 'dates'
        rows = cursor.fetchall()
        cursor.close()
        connection.close()

        # Convert dates to string format '%Y-%m-%d' for proper rendering
        holidays = [{"id": row[0], "dates": row[1].strftime('%d-%m-%Y')} for row in rows]
        return jsonify({"holidays": holidays})
    
    except psycopg2.Error as err:
        return jsonify({"error": str(err)}), 500

    
@app.route('/add_holiday', methods=['POST'])
def add_holiday():
    data = request.get_json()
    table = data.get('table', 'general_holidays')  # Default to 'general_holidays'
    holiday_date = data.get('date')

    try:
        connection = psycopg2.connect(**DB_CONFIG)
        cursor = connection.cursor()
        cursor.execute(f"INSERT INTO {table} (dates) VALUES (%s) RETURNING id", (holiday_date,))
        holiday_id = cursor.fetchone()[0]
        connection.commit()
        cursor.close()
        connection.close()
        return jsonify({"id": holiday_id, "date": holiday_date})
    except psycopg2.Error as e:
        return jsonify({"error": str(e)}), 500
    

@app.route('/delete_holiday', methods=['POST'])
def delete_holiday():
    data = request.get_json()
    table = data.get('table', 'general_holidays')  # Default to 'general_holidays'
    holiday_id = data.get('holiday_id')

    try:
        connection = psycopg2.connect(**DB_CONFIG)
        cursor = connection.cursor()
        cursor.execute(f"DELETE FROM {table} WHERE id = %s", (holiday_id,))
        connection.commit()
        cursor.close()
        connection.close()
        return jsonify({"success": True})
    except psycopg2.Error as e:
        return jsonify({"error": str(e)}), 500
    
from datetime import datetime
@app.route('/get_filtered_dates', methods=['POST'])
def get_filtered_dates():
    data = request.get_json()
    print("Received data:", data)  # 👈 DEBUG print

    year = data.get('year')
    semi_final_dates = data.get('semiFinalDates', [])

    if not year:
        return jsonify({"error": "Year not provided"}), 400

    if not semi_final_dates:
        return jsonify({"filteredDates": []})

    # Convert to YYYY-MM-DD format
    semi_final_dates_converted = []
    for date in semi_final_dates:
        try:
            converted_date = datetime.strptime(date, '%d/%m/%Y').strftime('%Y-%m-%d')
            semi_final_dates_converted.append(converted_date)
        except ValueError:
            print(f"Skipping invalid date: {date}")  # 👈 DEBUG print
            continue

    print("Converted dates:", semi_final_dates_converted)  # 👈 DEBUG print

    table_map = {
        'FE': 'fe',
        'SE': 'se',
        'TE': 'te',
        'BE': 'be'
    }

    selected_table = table_map.get(year)
    if not selected_table:
        return jsonify({"error": "Invalid year"}), 400

    tables_to_check = ['general_holidays', selected_table]

    try:
        connection = psycopg2.connect(**DB_CONFIG)
        cursor = connection.cursor()

        existing_holidays = set()
        for table in tables_to_check:
            if semi_final_dates_converted:
                query = f"SELECT dates FROM {table} WHERE dates IN %s"
                cursor.execute(query, (tuple(semi_final_dates_converted),))
                rows = cursor.fetchall()
                for row in rows:
                    existing_holidays.add(row[0].strftime('%Y-%m-%d'))

        cursor.close()
        connection.close()

        filtered_dates = [d for d in semi_final_dates_converted if d not in existing_holidays]
        return jsonify({"filteredDates": filtered_dates})

    except Exception as err:
        print("❌ ERROR:", str(err))  # 👈 LOG actual error
        return jsonify({"error": str(err)}), 500



# @app.route('/get_filtered_dates', methods=['POST'])
# def get_filtered_dates():
#     data = request.get_json()
#     year = data.get('year')
#     semi_final_dates = data.get('semiFinalDates', [])

#     if not semi_final_dates:
#         return jsonify({"filteredDates": []})

#     semi_final_dates_converted = []
#     for date in semi_final_dates:
#         try:
#             converted_date = datetime.strptime(date, '%d/%m/%Y').strftime('%Y-%m-%d')
#             semi_final_dates_converted.append(converted_date)
#         except ValueError:
#             continue

#     # Check individual tables for each year
#     if year == 'FE':
#         tables_to_check = ['general_holidays', 'fe']
#     elif year == 'SE':
#         tables_to_check = ['general_holidays', 'se']
#     elif year == 'TE':
#         tables_to_check = ['general_holidays', 'te']
#     elif year == 'BE':
#         tables_to_check = ['general_holidays', 'be']
#     else:
#         return jsonify({"error": "Invalid Year. Please enter FE, SE, TE, or BE."}), 400

#     existing_holidays = set()

#     try:
#         connection = psycopg2.connect(**DB_CONFIG)
#         cursor = connection.cursor()

#         for table in tables_to_check:
#             if semi_final_dates_converted:
#                 query = f"SELECT dates FROM {table} WHERE dates IN %s"
#                 cursor.execute(query, (tuple(semi_final_dates_converted),))
#                 rows = cursor.fetchall()
#                 existing_holidays.update([row[0].strftime('%Y-%m-%d') for row in rows])

#         cursor.close()
#         connection.close()

#         filtered_dates = [d for d in semi_final_dates_converted if d not in existing_holidays]
#         return jsonify({"filteredDates": filtered_dates})

#     except psycopg2.Error as err:
#         return jsonify({"error": str(err)}), 500


def filter_dates(start_date, end_date, holidays):
    from datetime import datetime, timedelta

    start_date_obj = datetime.strptime(start_date, '%Y-%m-%d')
    end_date_obj = datetime.strptime(end_date, '%Y-%m-%d')

    # Generate a list of all dates in the given range
    all_dates = []
    current_date = start_date_obj
    while current_date <= end_date_obj:
        all_dates.append(current_date.strftime('%Y-%m-%d'))
        current_date += timedelta(days=1)

    # Exclude holidays from the range
    filtered_dates = [date for date in all_dates if date not in holidays]
    return filtered_dates


    
@app.route('/files/<path:filename>')
def serve_file(filename):
    return send_from_directory('files', filename)



@app.route('/')
def index():
    return render_template('index.html')

@app.route('/regular')
def regular():
    return render_template('regular.html')

@app.route('/lab')
def lab():
    return render_template('lab.html')


@app.route('/process', methods=['POST'])
def process_document():
    if 'file' not in request.files:
        return 'No file part'

    file = request.files['file']

    if file.filename == '':
        return 'No selected file'

    if file:
        document = Document(file)
        content = "<table border='1' id='resultTable'>"

        for table in document.tables:
            for row in table.rows:
                content += "<tr>"
                for cell in row.cells:
                    content += f"<td>{cell.text}</td>"
                content += "</tr>"

        content += "</table>"
        return content
    



def is_bold(run):
    return run.bold



def generate_table(header, rows):
    table_html = "<table border='1'>"
    table_html += "<tr>"

    
    for i, col in enumerate(header):
        if i != 1 and i != 3:
            table_html += f"<th>{col}</th>"
    table_html += "</tr>"

    if rows:
        
        try:
            k = int(header[-1])
            if k > 0:
                cell_contents = [[] for _ in range(len(header) - 2)]  

                for i, row in enumerate(rows):
                    for j, cell in enumerate(row[:1] + row[2:3] + row[4:]): 
                       
                        cell_content = cell.split(',')
                        cell_contents[j].extend(cell_content)

                items_per_cell = max(len(content) // k + (1 if len(content) % k > 0 else 0) for content in cell_contents)

                for i in range(k):
                    table_html += "<tr>"
                    for j in range(len(header) - 2):  
                        start_idx = i * items_per_cell
                        end_idx = min((i + 1) * items_per_cell, len(cell_contents[j]))
                        table_html += f"<td>{', '.join(cell_contents[j][start_idx:end_idx])}</td>"
                    table_html += "</tr>"


        except ValueError:
            pass  

    table_html += "</table>"
    return table_html





@app.route('/process_A', methods=['POST'])
def process_document_A():
    if 'file' not in request.files:
        return 'No file part'

    file = request.files['file']

    if file.filename == '':
        return 'No selected file'

    if file:
        document = Document(file)
        content = ""
        current_table_header = []
        current_table_rows = []
        last_column_values = []  

        table_count = 0  

        for table in document.tables:
            table_count += 1
            if table_count < 3:
                continue  
            elif table_count > 4:
                break 

            for row in table.rows:
     
                last_cell_text = row.cells[-1].text.strip()
                last_column_values.append(last_cell_text)

                first_column_text = row.cells[0].text.strip()
                if first_column_text:
                    if current_table_header and current_table_rows:
                        content += generate_table(current_table_header, current_table_rows)


                    current_table_header = [cell.text.strip() for cell in row.cells]
                    current_table_rows = []
                else:
                    current_table_rows.append([cell.text.strip() for cell in row.cells])

        if current_table_header and current_table_rows:
            content += generate_table(current_table_header, current_table_rows)

        if content == "":
            return 'No tables found'
        else:
            response_data = {
                'content': content,
                'last_column_values': last_column_values
            }
            return jsonify(response_data)




@app.route('/process_B', methods=['POST'])
def process_document_B():
    if 'file' not in request.files:
        return 'No file part'

    file = request.files['file']

    if file.filename == '':
        return 'No selected file'

    if file:
        document = Document(file)
        content = ""

        table_count = 0
        for table in document.tables:
            table_count += 1
            if table_count == 3:
                content += "<table border='1'>"

                for row in table.rows:
                    first_cell_text = row.cells[0].text.strip()
                    if first_cell_text.isdigit():  # Check if first cell contains a natural number
                        content += "<tr>"
                        content += f"<td>{row.cells[1].text.strip()}</td>"
                        content += "</tr>"

                content += "</table>"
                break

        if content == "<table border='1'></table>":
            return 'No valid rows found in the third table'
        else:
            return content











# if __name__ == '__main__':
#     os.makedirs(FILES_DIR, exist_ok=True)
#     app.run(debug=True, host='0.0.0.0')


if __name__ == '__main__':
   os.makedirs(FILES_DIR, exist_ok=True)
   app.run(debug=True)



